"""Document ingestion service."""
import os
import uuid
from typing import List, Optional, Dict, Any
from docx import Document
from sqlalchemy.orm import Session

from app.models.toolkit import ToolkitDocument, ToolkitChunk


def parse_docx(file_path: str) -> List[Dict[str, Any]]:
    """
    Parse DOCX file and extract structured content.

    Args:
        file_path: Path to DOCX file

    Returns:
        List of content blocks with text and metadata
    """
    doc = Document(file_path)
    content_blocks = []
    current_heading = None

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Check if paragraph is a heading
        if para.style.name.startswith('Heading'):
            current_heading = text
            content_blocks.append({
                'type': 'heading',
                'text': text,
                'level': para.style.name,
                'heading': current_heading
            })
        else:
            content_blocks.append({
                'type': 'paragraph',
                'text': text,
                'heading': current_heading
            })

    return content_blocks


def chunk_content(
    content_blocks: List[Dict[str, Any]],
    target_size: int = 1000,
    overlap: int = 150
) -> List[Dict[str, Any]]:
    """
    Chunk content into manageable pieces with overlap.

    Args:
        content_blocks: List of content blocks from parse_docx
        target_size: Target chunk size in characters (800-1200)
        overlap: Overlap between chunks in characters

    Returns:
        List of chunks with text and metadata
    """
    chunks = []
    current_chunk = []
    current_size = 0
    current_heading = None
    chunk_index = 0

    for block in content_blocks:
        if block['type'] == 'heading':
            current_heading = block['text']

        text = block['text']
        text_len = len(text)

        # If adding this block would exceed target size and we have content, create chunk
        if current_size + text_len > target_size and current_chunk:
            chunk_text = ' '.join(current_chunk)
            chunks.append({
                'chunk_text': chunk_text,
                'chunk_index': chunk_index,
                'heading': current_heading,
                'metadata': {'char_count': len(chunk_text)}
            })
            chunk_index += 1

            # Keep last part for overlap
            if current_chunk:
                overlap_text = current_chunk[-1]
                if len(overlap_text) > overlap:
                    overlap_text = overlap_text[-overlap:]
                current_chunk = [overlap_text, text]
                current_size = len(overlap_text) + text_len
            else:
                current_chunk = [text]
                current_size = text_len
        else:
            current_chunk.append(text)
            current_size += text_len

    # Add remaining content as final chunk
    if current_chunk:
        chunk_text = ' '.join(current_chunk)
        chunks.append({
            'chunk_text': chunk_text,
            'chunk_index': chunk_index,
            'heading': current_heading,
            'metadata': {'char_count': len(chunk_text)}
        })

    return chunks


def ingest_document(
    db: Session,
    file_path: str,
    version_tag: str,
    source_filename: str,
    create_embeddings: bool = True
) -> ToolkitDocument:
    """
    Ingest a document: parse, chunk, and store in database.

    Args:
        db: Database session
        file_path: Path to uploaded DOCX file
        version_tag: Version identifier
        source_filename: Original filename
        create_embeddings: Whether to create embeddings (requires OpenAI API)

    Returns:
        Created ToolkitDocument instance
    """
    # Check if version already exists
    existing = db.query(ToolkitDocument).filter(
        ToolkitDocument.version_tag == version_tag
    ).first()
    if existing:
        raise ValueError(f"Version tag '{version_tag}' already exists")

    # Parse document
    content_blocks = parse_docx(file_path)

    # Create chunks
    chunks = chunk_content(content_blocks)

    # Create document record
    doc = ToolkitDocument(
        version_tag=version_tag,
        source_filename=source_filename,
        file_path=file_path,
        chunk_count=len(chunks)
    )
    db.add(doc)
    db.flush()  # Get document ID

    # Create chunk records
    chunk_objects = []
    for chunk_data in chunks:
        chunk = ToolkitChunk(
            document_id=doc.id,
            chunk_text=chunk_data['chunk_text'],
            chunk_index=chunk_data['chunk_index'],
            heading=chunk_data.get('heading'),
            chunk_metadata=chunk_data.get('metadata'),
            embedding=None  # Will be populated by embeddings service
        )
        chunk_objects.append(chunk)

    db.bulk_save_objects(chunk_objects)
    db.commit()
    db.refresh(doc)

    # Create embeddings if requested
    if create_embeddings:
        from app.services.embeddings import create_embeddings_for_document
        create_embeddings_for_document(db, doc.id)

    return doc


def reindex_document(db: Session, document_id: str) -> ToolkitDocument:
    """
    Reindex a document: re-run chunking and embeddings.

    This deletes existing chunks and recreates them from the source file.

    Args:
        db: Database session
        document_id: ID of document to reindex

    Returns:
        Updated ToolkitDocument instance

    Raises:
        ValueError: If document not found or file doesn't exist
    """
    # Get document
    doc = db.query(ToolkitDocument).filter(ToolkitDocument.id == document_id).first()

    if not doc:
        raise ValueError(f"Document {document_id} not found")

    if not os.path.exists(doc.file_path):
        raise ValueError(f"Source file not found: {doc.file_path}")

    # Delete existing chunks
    db.query(ToolkitChunk).filter(ToolkitChunk.document_id == document_id).delete()
    db.commit()

    # Re-parse document
    content_blocks = parse_docx(doc.file_path)

    # Create new chunks
    chunks = chunk_content(content_blocks)

    # Update chunk count
    doc.chunk_count = len(chunks)

    # Create chunk records
    chunk_objects = []
    for chunk_data in chunks:
        chunk = ToolkitChunk(
            document_id=doc.id,
            chunk_text=chunk_data['chunk_text'],
            chunk_index=chunk_data['chunk_index'],
            heading=chunk_data.get('heading'),
            chunk_metadata=chunk_data.get('metadata'),
            embedding=None  # Will be populated by embeddings service
        )
        chunk_objects.append(chunk)

    db.bulk_save_objects(chunk_objects)
    db.commit()
    db.refresh(doc)

    # Recreate embeddings
    from app.services.embeddings import create_embeddings_for_document
    create_embeddings_for_document(db, doc.id)

    return doc
