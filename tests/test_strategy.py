"""Strategy plan tests."""
import pytest
import os
import tempfile
from docx import Document
from sqlalchemy.orm import Session

from app.models.auth import User
from app.models.toolkit import StrategyPlan
from app.services.auth import create_user
from app.services.ingestion import ingest_document
from app.services.strategy import generate_strategy_plan, export_plan_to_markdown


def create_test_user(db: Session, username: str = "testuser", email: str = "test@example.com") -> User:
    """Create a test user."""
    return create_user(db, email=email, username=username, password="password123")


def create_test_document() -> str:
    """Create a test DOCX file with strategy-relevant content."""
    doc = Document()
    doc.add_heading('AI Tool Recommendations', level=1)
    doc.add_paragraph('For code assistance, we recommend using GitHub Copilot or ChatGPT for improved productivity.')

    doc.add_heading('Security Best Practices', level=1)
    doc.add_paragraph('Always validate AI outputs before using in production. Implement proper access controls and data privacy measures.')

    doc.add_heading('Cloud Deployment', level=1)
    doc.add_paragraph('Cloud-based AI tools offer scalability and ease of deployment. Consider data residency requirements.')

    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    temp_file.close()
    return temp_file.name


def test_create_strategy_plan_persists(db_session, monkeypatch):
    """Test that strategy plan is created and persisted to database."""
    from app import settings as app_settings
    monkeypatch.setattr(app_settings.settings, "EMBEDDING_PROVIDER", "local_stub")

    # Create user
    user = create_test_user(db_session)

    # Ingest toolkit content
    docx_file = create_test_document()
    try:
        ingest_document(
            db=db_session,
            file_path=docx_file,
            version_tag="strategy-test-v1",
            source_filename="strategy-test.docx",
            create_embeddings=True
        )
    finally:
        os.unlink(docx_file)

    # Create inputs
    inputs = {
        "role": "CTO",
        "org_type": "startup",
        "risk_level": "medium",
        "data_sensitivity": "internal",
        "budget": "small",
        "deployment_pref": "cloud",
        "use_cases": ["code_assistance", "content_generation"]
    }

    # Generate strategy plan
    # Note: This will try to call OpenAI, so we need to mock it
    from app.services import strategy
    from unittest.mock import Mock

    # Mock OpenAI call
    original_generate = strategy._generate_grounded_plan

    def mock_generate(inputs, chunks):
        plan_text = "# Test Strategy Plan\n\nBased on the toolkit content [1], we recommend using AI tools for code assistance."
        citations = [
            {
                "chunk_id": chunks[0].chunk_id if chunks else "test-id",
                "heading": chunks[0].heading if chunks else "Test",
                "excerpt": chunks[0].chunk_text[:200] if chunks else "Test excerpt",
                "similarity_score": 0.85,
                "cluster": None,
                "tool_name": None
            }
        ] if chunks else []
        return plan_text, citations

    monkeypatch.setattr(strategy, "_generate_grounded_plan", mock_generate)

    # Generate plan
    plan = generate_strategy_plan(
        db=db_session,
        user_id=str(user.id),
        inputs=inputs
    )

    # Verify plan was created
    assert plan is not None
    assert plan.id is not None
    assert plan.user_id == user.id

    # Verify plan persisted
    saved_plan = db_session.query(StrategyPlan).filter(
        StrategyPlan.id == plan.id
    ).first()

    assert saved_plan is not None
    assert saved_plan.user_id == user.id
    assert saved_plan.inputs == inputs
    assert len(saved_plan.plan_text) > 0
    assert isinstance(saved_plan.citations, list)


def test_strategy_plan_belongs_to_user(db_session, monkeypatch):
    """Test that strategy plan belongs to the user who created it."""
    from app import settings as app_settings
    monkeypatch.setattr(app_settings.settings, "EMBEDDING_PROVIDER", "local_stub")

    # Create two users
    user1 = create_test_user(db_session, username="user1", email="user1@example.com")
    user2 = create_test_user(db_session, username="user2", email="user2@example.com")

    # Ingest content
    docx_file = create_test_document()
    try:
        ingest_document(
            db=db_session,
            file_path=docx_file,
            version_tag="strategy-isolation-test-v1",
            source_filename="strategy-isolation-test.docx",
            create_embeddings=True
        )
    finally:
        os.unlink(docx_file)

    # Mock OpenAI
    from app.services import strategy

    def mock_generate(inputs, chunks):
        return "Test plan", []

    monkeypatch.setattr(strategy, "_generate_grounded_plan", mock_generate)

    # Create plan for user1
    inputs = {
        "role": "Developer",
        "org_type": "startup",
        "risk_level": "high",
        "data_sensitivity": "public",
        "budget": "minimal",
        "deployment_pref": "cloud",
        "use_cases": ["code_assistance"]
    }

    plan = generate_strategy_plan(
        db=db_session,
        user_id=str(user1.id),
        inputs=inputs
    )

    # Verify user1 can access their plan
    user1_plan = db_session.query(StrategyPlan).filter(
        StrategyPlan.id == plan.id,
        StrategyPlan.user_id == user1.id
    ).first()

    assert user1_plan is not None

    # Verify user2 cannot access user1's plan
    user2_plan = db_session.query(StrategyPlan).filter(
        StrategyPlan.id == plan.id,
        StrategyPlan.user_id == user2.id
    ).first()

    assert user2_plan is None


def test_strategy_plan_not_accessible_to_other_users(db_session, monkeypatch):
    """Test that users cannot access other users' strategy plans."""
    from app import settings as app_settings
    monkeypatch.setattr(app_settings.settings, "EMBEDDING_PROVIDER", "local_stub")

    # Create two users
    user1 = create_test_user(db_session, username="alice", email="alice@example.com")
    user2 = create_test_user(db_session, username="bob", email="bob@example.com")

    # Ingest content
    docx_file = create_test_document()
    try:
        ingest_document(
            db=db_session,
            file_path=docx_file,
            version_tag="strategy-access-test-v1",
            source_filename="strategy-access-test.docx",
            create_embeddings=True
        )
    finally:
        os.unlink(docx_file)

    # Mock OpenAI
    from app.services import strategy

    def mock_generate(inputs, chunks):
        return "Test plan for " + inputs['role'], []

    monkeypatch.setattr(strategy, "_generate_grounded_plan", mock_generate)

    # Create plan for user1
    inputs1 = {
        "role": "CTO",
        "org_type": "enterprise",
        "risk_level": "low",
        "data_sensitivity": "regulated",
        "budget": "large",
        "deployment_pref": "sovereign",
        "use_cases": ["data_analysis"]
    }

    plan1 = generate_strategy_plan(
        db=db_session,
        user_id=str(user1.id),
        inputs=inputs1
    )

    # Create plan for user2
    inputs2 = {
        "role": "Developer",
        "org_type": "startup",
        "risk_level": "high",
        "data_sensitivity": "public",
        "budget": "minimal",
        "deployment_pref": "cloud",
        "use_cases": ["code_assistance"]
    }

    plan2 = generate_strategy_plan(
        db=db_session,
        user_id=str(user2.id),
        inputs=inputs2
    )

    # Get all plans for user1
    user1_plans = db_session.query(StrategyPlan).filter(
        StrategyPlan.user_id == user1.id
    ).all()

    # Get all plans for user2
    user2_plans = db_session.query(StrategyPlan).filter(
        StrategyPlan.user_id == user2.id
    ).all()

    # User1 should only see their own plan
    assert len(user1_plans) == 1
    assert user1_plans[0].id == plan1.id

    # User2 should only see their own plan
    assert len(user2_plans) == 1
    assert user2_plans[0].id == plan2.id

    # Plans should be different
    assert plan1.id != plan2.id


def test_export_plan_to_markdown(db_session, monkeypatch):
    """Test exporting strategy plan to Markdown format."""
    from app import settings as app_settings
    monkeypatch.setattr(app_settings.settings, "EMBEDDING_PROVIDER", "local_stub")

    # Create user
    user = create_test_user(db_session)

    # Ingest content
    docx_file = create_test_document()
    try:
        ingest_document(
            db=db_session,
            file_path=docx_file,
            version_tag="strategy-export-test-v1",
            source_filename="strategy-export-test.docx",
            create_embeddings=True
        )
    finally:
        os.unlink(docx_file)

    # Mock OpenAI
    from app.services import strategy

    def mock_generate(inputs, chunks):
        plan_text = "# Strategy Plan\n\nRecommendations based on toolkit [1]."
        citations = [
            {
                "chunk_id": "test-id",
                "heading": "Best Practices",
                "excerpt": "Test content excerpt",
                "similarity_score": 0.9,
                "cluster": "practices",
                "tool_name": "TestTool"
            }
        ]
        return plan_text, citations

    monkeypatch.setattr(strategy, "_generate_grounded_plan", mock_generate)

    # Create plan
    inputs = {
        "role": "PM",
        "org_type": "sme",
        "risk_level": "medium",
        "data_sensitivity": "internal",
        "budget": "medium",
        "deployment_pref": "hybrid",
        "use_cases": ["automation"]
    }

    plan = generate_strategy_plan(
        db=db_session,
        user_id=str(user.id),
        inputs=inputs
    )

    # Export to Markdown
    markdown = export_plan_to_markdown(plan)

    # Verify Markdown content
    assert "# AI Implementation Strategy Plan" in markdown
    assert "Role" in markdown
    assert "PM" in markdown
    assert "Strategy Plan" in markdown
    assert "Sources & Citations" in markdown
    assert "Best Practices" in markdown


def test_strategy_plan_includes_citations(db_session, monkeypatch):
    """Test that strategy plan includes citations from toolkit chunks."""
    from app import settings as app_settings
    monkeypatch.setattr(app_settings.settings, "EMBEDDING_PROVIDER", "local_stub")

    # Create user
    user = create_test_user(db_session)

    # Ingest content
    docx_file = create_test_document()
    try:
        ingest_document(
            db=db_session,
            file_path=docx_file,
            version_tag="strategy-citations-test-v1",
            source_filename="strategy-citations-test.docx",
            create_embeddings=True
        )
    finally:
        os.unlink(docx_file)

    # Mock OpenAI
    from app.services import strategy

    def mock_generate(inputs, chunks):
        citations = []
        for chunk in chunks[:3]:
            citations.append({
                "chunk_id": chunk.chunk_id,
                "heading": chunk.heading,
                "excerpt": chunk.chunk_text[:200],
                "similarity_score": chunk.similarity_score,
                "cluster": chunk.cluster,
                "tool_name": chunk.tool_name
            })
        return "Plan with citations [1], [2], [3]", citations

    monkeypatch.setattr(strategy, "_generate_grounded_plan", mock_generate)

    # Create plan
    inputs = {
        "role": "Architect",
        "org_type": "enterprise",
        "risk_level": "low",
        "data_sensitivity": "pii",
        "budget": "large",
        "deployment_pref": "sovereign",
        "use_cases": ["research", "automation"]
    }

    plan = generate_strategy_plan(
        db=db_session,
        user_id=str(user.id),
        inputs=inputs
    )

    # Verify citations exist
    assert plan.citations is not None
    assert isinstance(plan.citations, list)
    assert len(plan.citations) > 0

    # Each citation should have required fields
    for citation in plan.citations:
        assert "chunk_id" in citation
        assert "heading" in citation
        assert "excerpt" in citation
